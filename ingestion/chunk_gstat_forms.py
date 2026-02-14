import os
import re
import json
import uuid
from docx import Document

# ==========================
# CONFIG
# ==========================

FORMS_FOLDER = "data/raw/docx/gstat_forms"
OUTPUT_FILE = "data/processed/gstat_forms.json"

PARENT_DOC = "GSTAT Rules 2025"
HIERARCHY_LEVEL = 4

# ==========================
# LOAD DOCX
# ==========================

def load_docx(file_path):
    return Document(file_path)

def load_docx_text(doc):
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())
    return "\n".join(text)

# ==========================
# EXTRACT DOCUMENT METADATA
# ==========================

def extract_document_metadata(text):

    # FORM detection
    form_match = re.search(r"GSTAT\s*FORM[-\s]*(\d+)", text, re.IGNORECASE)

    # CDR detection
    cdr_match = re.search(r"GSTAT[-\s]*CDR[-\s]*(\d+)", text, re.IGNORECASE)

    # Extract title after dash
    title_match = re.search(r"–\s*(.*?)\n", text)
    title = title_match.group(1).strip() if title_match else "Untitled"

    # Extract rule reference (handles 59(c))
    rule_match = re.search(r"\[See\s*rule\s*([\d\(\)a-zA-Z]+)\]", text, re.IGNORECASE)
    related_rule = rule_match.group(1) if rule_match else None

    if form_match:
        return {
            "doc_type": "Form",
            "number": form_match.group(1),
            "title": title,
            "related_rule": related_rule
        }

    if cdr_match:
        return {
            "doc_type": "Register",
            "number": cdr_match.group(1),
            "title": title,
            "related_rule": related_rule
        }

    return {
        "doc_type": "Other",
        "number": "UNKNOWN",
        "title": title,
        "related_rule": related_rule
    }

# ==========================
# EXTRACT FORM FIELDS
# ==========================

def extract_form_fields(text):
    matches = re.findall(r"\d+\.\s*(.*?)\s*–", text)
    return [m.strip() for m in matches]

# ==========================
# EXTRACT TABLE HEADERS (Important for CDR)
# ==========================

def extract_table_headers(doc):
    headers = []

    for table in doc.tables:
        if table.rows:
            first_row = table.rows[0]
            for cell in first_row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    headers.append(cell_text)

    return headers

# ==========================
# GENERATE SUMMARY
# ==========================

def generate_summary(doc_type, number, title):
    return f"GSTAT {doc_type} {number} relates to {title}."

# ==========================
# PROCESS SINGLE FILE
# ==========================

def process_file(file_path):

    doc = load_docx(file_path)
    text = load_docx_text(doc)

    meta = extract_document_metadata(text)

    doc_type = meta["doc_type"]
    doc_number = meta["number"]
    doc_title = meta["title"]
    related_rule = meta["related_rule"]

    form_fields = extract_form_fields(text)
    table_headers = extract_table_headers(doc)

    summary = generate_summary(doc_type, doc_number, doc_title)

    # Determine chunk_type for legal hierarchy
    if doc_type == "Form":
        chunk_type = "gstat_form"
    elif doc_type == "Register":
        chunk_type = "gstat_register"
    else:
        chunk_type = "gstat_form"  # Default for "Other" types related to GSTAT

    base_chunk = {
        "id": str(uuid.uuid4()),
        "doc_type": doc_type,
        "chunk_type": chunk_type,  # For legal hierarchy prioritization
        "parent_doc": PARENT_DOC,
        "hierarchy_level": HIERARCHY_LEVEL,

        "structure": {
            "number": doc_number,
            "title": doc_title,
            "related_rule": related_rule
        },

        "text": text,
        "summary": summary,
        "keywords": form_fields[:5] if form_fields else table_headers[:5],

        "metadata": {
            "source": "GSTAT Forms",
            "source_file": os.path.basename(file_path)
        }
    }

    # Attach structured fields properly
    if doc_type == "Form":
        base_chunk["form_fields"] = form_fields

    if doc_type == "Register":
        base_chunk["register_columns"] = table_headers

    chunks = [base_chunk]

    # ==========================
    # HANDLE SCHEDULE OF FEES
    # ==========================

    if "SCHEDULE OF FEES" in text.upper():

        fee_section = text.upper().split("SCHEDULE OF FEES")[-1]

        fee_chunk = {
            "id": str(uuid.uuid4()),
            "doc_type": "Schedule",
            "chunk_type": "gstat_form",  # Fee schedules also part of GSTAT forms
            "parent_doc": PARENT_DOC,
            "hierarchy_level": HIERARCHY_LEVEL,

            "structure": {
                "related_document": doc_number,
                "type": "Fee Schedule"
            },

            "text": fee_section.strip(),
            "summary": f"Fee schedule applicable under GSTAT {doc_type} {doc_number}.",
            "keywords": ["fees", "schedule", "application", "amount"],

            "metadata": {
                "source_file": os.path.basename(file_path)
            }
        }

        chunks.append(fee_chunk)

    return chunks

# ==========================
# MAIN
# ==========================

def main():

    all_chunks = []

    for filename in os.listdir(FORMS_FOLDER):

        if filename.lower().endswith(".docx"):

            file_path = os.path.join(FORMS_FOLDER, filename)
            print(f"Processing {filename}...")

            file_chunks = process_file(file_path)
            all_chunks.extend(file_chunks)

    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(all_chunks, f, indent=4, ensure_ascii=False)

    print(f"\nCreated {len(all_chunks)} chunks from forms folder.")

if __name__ == "__main__":
    main()
