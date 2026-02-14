import os
import re
import json
import uuid
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
import nltk

nltk.download('punkt')

# ==========================
# CONFIG
# ==========================

INPUT_FOLDER = "data/raw/docx/case_studies"
OUTPUT_FILE = "data/processed/case_studies.json"

HIERARCHY_LEVEL = 5
DOC_TYPE = "Case Study"

# ==========================
# LOAD DOCX
# ==========================

def load_docx(file_path):

    doc = Document(file_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    return paragraphs, tables

# ==========================
# SPLIT CASE SECTIONS
# ==========================

def split_sections(paragraphs):

    text = "\n".join(paragraphs)

    sections = []

    # Split Questions
    question_pattern = r'(?:Q\.?|Question)\s*\d+.*?(?=(?:Q\.?|Question)\s*\d+|$)'
    question_blocks = re.findall(question_pattern, text, re.DOTALL | re.IGNORECASE)

    # Split Answers
    answer_pattern = r'(?:Ans\.?|Answer)\s*\d*.*?(?=(?:Ans\.?|Answer)\s*\d*|$)'
    answer_blocks = re.findall(answer_pattern, text, re.DOTALL | re.IGNORECASE)

    # Facts = before first Question
    first_question_index = re.search(r'(?:Q\.?|Question)\s*\d+', text)
    if first_question_index:
        facts = text[:first_question_index.start()]
        sections.append(("Facts", None, facts.strip()))

    # Pair Questions & Answers
    for i in range(len(question_blocks)):
        q_text = question_blocks[i].strip()
        a_text = answer_blocks[i].strip() if i < len(answer_blocks) else ""

        q_num_match = re.search(r'\d+', q_text)
        q_number = q_num_match.group(0) if q_num_match else str(i+1)

        sections.append(("Question", q_number, q_text))
        sections.append(("Answer", q_number, a_text))

    return sections

# ==========================
# SUMMARY
# ==========================

def generate_summary(text):
    sentences = nltk.sent_tokenize(text)
    return sentences[0] if sentences else text[:200]

# ==========================
# KEYWORDS
# ==========================

def extract_keywords(text, top_n=8):

    try:
        vectorizer = TfidfVectorizer(
            stop_words='english',
            max_features=200,
            ngram_range=(1,2)
        )

        tfidf_matrix = vectorizer.fit_transform([text])
        scores = zip(vectorizer.get_feature_names_out(), tfidf_matrix.toarray()[0])
        sorted_keywords = sorted(scores, key=lambda x: x[1], reverse=True)

        return [word for word, score in sorted_keywords[:top_n]]

    except:
        return []

# ==========================
# PROCESS SINGLE FILE
# ==========================

def process_case_study(file_path):

    print(f"Processing {os.path.basename(file_path)}")

    paragraphs, tables = load_docx(file_path)
    sections = split_sections(paragraphs)

    case_name = os.path.splitext(os.path.basename(file_path))[0]

    chunks = []

    for section_type, q_number, text in sections:

        chunk = {
            "id": str(uuid.uuid4()),
            "doc_type": DOC_TYPE,
            "parent_doc": case_name,
            "hierarchy_level": HIERARCHY_LEVEL,

            "structure": {
                "section_type": section_type,
                "question_number": q_number
            },

            "text": text,
            "summary": generate_summary(text),
            "keywords": extract_keywords(text),

            "metadata": {
                "source": "Case Studies",
                "source_file": os.path.basename(file_path)
            }
        }

        chunks.append(chunk)

    # Add Tables separately
    for idx, table in enumerate(tables):

        chunk = {
            "id": str(uuid.uuid4()),
            "doc_type": "Case Study Table",
            "parent_doc": case_name,
            "hierarchy_level": HIERARCHY_LEVEL,

            "structure": {
                "table_index": idx + 1
            },

            "text": str(table),
            "summary": f"Financial table {idx+1} in case study.",
            "keywords": [],

            "metadata": {
                "source_file": os.path.basename(file_path)
            }
        }

        chunks.append(chunk)

    return chunks

# ==========================
# MAIN
# ==========================

def main():

    all_chunks = []

    for filename in os.listdir(INPUT_FOLDER):

        if filename.lower().endswith(".docx"):
            file_path = os.path.join(INPUT_FOLDER, filename)
            chunks = process_case_study(file_path)
            all_chunks.extend(chunks)

    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(all_chunks, f, indent=4, ensure_ascii=False)

    print(f"\nCreated {len(all_chunks)} case study chunks.")

if __name__ == "__main__":
    main()
